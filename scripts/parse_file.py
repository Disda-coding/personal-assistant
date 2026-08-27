#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""通用文件解析脚本（含离线降级，纯标准库实现降级路径）。

用法:
    python parse_file.py <文件路径> [--max-chars 50000]

支持: 纯文本类(txt/md/代码/json/csv 等), docx, xlsx, pptx, pdf
二进制格式优先使用已安装的第三方库; 库缺失时降级为 zipfile+正则 的零依赖方案;
pdf 在无 pypdf 时输出安装引导而不是报错崩溃（离线环境友好）。
"""
import argparse
import os
import re
import sys
import zipfile


def xml_unescape(s):
    """解码 XML 转义: 命名实体 + 数字字符引用(&#NNN; / &#xHH;)。

    xml.sax.saxutils.unescape 不支持数字字符引用(如 openpyxl 生成的
    sharedStrings.xml 中会用到), 故自行实现。
    """
    table = {"amp": "&", "lt": "<", "gt": ">", "quot": '"', "apos": "'"}

    def repl(m):
        e = m.group(1)
        try:
            if e.startswith(("#x", "#X")):
                return chr(int(e[2:], 16))
            if e.startswith("#"):
                return chr(int(e[1:]))
        except (ValueError, OverflowError):
            return m.group(0)
        return table.get(e, m.group(0))

    return re.sub(r"&(#x[0-9a-fA-F]+|#[0-9]+|\w+);", repl, s)

TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".json", ".csv", ".tsv", ".xml", ".yaml", ".yml",
    ".html", ".htm", ".js", ".jsx", ".ts", ".tsx", ".py", ".java", ".c", ".cpp",
    ".h", ".hpp", ".cs", ".go", ".rs", ".rb", ".php", ".sh", ".bat", ".ps1",
    ".ini", ".cfg", ".conf", ".log", ".sql", ".css", ".scss", ".less", ".vue",
    ".toml", ".env",
}


def reconfigure_stdout():
    # 防止 Windows 控制台 GBK 乱码
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass


def read_text_file(path):
    for enc in ("utf-8-sig", "utf-8", "gbk"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            continue
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def is_binary(path):
    with open(path, "rb") as f:
        return b"\x00" in f.read(1024)


def _num_key(name):
    m = re.search(r"(\d+)", name)
    return int(m.group(1)) if m else 0


# ---------- docx ----------

def parse_docx(path):
    try:
        import docx  # python-docx
    except ImportError:
        return parse_docx_fallback(path)
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts), "python-docx"


def parse_docx_fallback(path):
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
    out = []
    for m in re.finditer(r"<w:t[^>]*>([^<]*)</w:t>|</w:p>", xml):
        out.append("\n" if m.group(1) is None else xml_unescape(m.group(1)))
    return "".join(out).strip(), "zipfile+正则 (降级, 零依赖)"


# ---------- xlsx ----------

def parse_xlsx(path):
    try:
        import openpyxl
    except ImportError:
        return parse_xlsx_fallback(path)
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"[工作表: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts), "openpyxl"


def parse_xlsx_fallback(path):
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        shared = []
        if "xl/sharedStrings.xml" in names:
            ss = z.read("xl/sharedStrings.xml").decode("utf-8", errors="ignore")
            shared = [xml_unescape(x) for x in re.findall(r"<t[^>]*>([^<]*)</t>", ss)]
        parts = []
        if shared:
            parts.append("[共享字符串]")
            parts.extend(s for s in shared if s.strip())
        sheets = sorted(
            (x for x in names if re.match(r"xl/worksheets/sheet\d+\.xml$", x)),
            key=_num_key,
        )
        for sn in sheets:
            parts.append(f"\n[{sn}]")
            xml = z.read(sn).decode("utf-8", errors="ignore")
            for rowm in re.finditer(r"<row[^>]*>(.*?)</row>", xml, re.S):
                cells = []
                for cm in re.finditer(r"<c\b[^>]*>(.*?)</c>", rowm.group(1), re.S):
                    opening_tag = cm.group(0)[: cm.start(1) - cm.start(0)]
                    inner = cm.group(1)
                    tm = re.search(r't="([^"]*)"', opening_tag)
                    ctype = tm.group(1) if tm else ""
                    vm = re.search(r"<v>([^<]*)</v>", inner)
                    im = re.search(r"<t[^>]*>([^<]*)</t>", inner)
                    if ctype == "s" and vm:
                        idx = int(vm.group(1))
                        cells.append(shared[idx] if idx < len(shared) else "")
                    elif im:
                        cells.append(xml_unescape(im.group(1)))
                    elif vm:
                        cells.append(vm.group(1))
                if cells:
                    parts.append(" | ".join(cells))
    return "\n".join(parts), "zipfile+正则 (降级, 零依赖)"


# ---------- pptx ----------

def parse_pptx(path):
    with zipfile.ZipFile(path) as z:
        slides = sorted(
            (x for x in z.namelist() if re.match(r"ppt/slides/slide\d+\.xml$", x)),
            key=_num_key,
        )
        parts = []
        for i, sn in enumerate(slides, 1):
            xml = z.read(sn).decode("utf-8", errors="ignore")
            texts = [xml_unescape(x) for x in re.findall(r"<a:t>([^<]*)</a:t>", xml)]
            texts = [t for t in texts if t.strip()]
            parts.append(f"--- 幻灯片 {i} ---")
            parts.append("\n".join(texts))
    return "\n".join(parts), "zipfile+正则 (零依赖)"


# ---------- pdf ----------

def parse_pdf(path):
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # 兼容旧包名
        except ImportError:
            return None
    reader = PdfReader(path)
    parts = []
    for i, page in enumerate(reader.pages, 1):
        t = (page.extract_text() or "").strip()
        parts.append(f"--- 第 {i} 页 ---\n{t}")
    return "\n".join(parts), "pypdf"


def print_pdf_guide():
    print("无法解析 PDF: 本机未安装 pypdf。")
    print()
    print("在线环境, 直接执行:")
    print("    pip install pypdf")
    print()
    print("离线环境, 在有网络的电脑上下载安装包:")
    print("    pip download pypdf -d ./pypkg")
    print("把 pypkg 文件夹拷贝到本机后执行:")
    print("    pip install pypdf --no-index --find-links ./pypkg")


def main():
    reconfigure_stdout()
    ap = argparse.ArgumentParser(description="通用文件解析（含离线降级）")
    ap.add_argument("file", help="文件路径")
    ap.add_argument("--max-chars", type=int, default=50000, help="输出最大字符数（默认 50000）")
    args = ap.parse_args()

    path = args.file
    if not os.path.isfile(path):
        print(f"错误: 文件不存在: {path}")
        return 1

    ext = os.path.splitext(path)[1].lower()

    try:
        if ext in TEXT_EXTS:
            text, method = read_text_file(path).strip(), "纯文本"
        elif ext == ".docx":
            text, method = parse_docx(path)
        elif ext == ".xlsx":
            text, method = parse_xlsx(path)
        elif ext == ".pptx":
            text, method = parse_pptx(path)
        elif ext == ".pdf":
            result = parse_pdf(path)
            if result is None:
                print_pdf_guide()
                return 3
            text, method = result
        else:
            if is_binary(path):
                print(f"无法解析: 不支持的二进制文件类型 {ext or '(无扩展名)'}")
                print("支持: 纯文本类(md/txt/代码/json/csv 等), docx, xlsx, pptx, pdf")
                return 2
            text, method = read_text_file(path).strip(), "纯文本(无扩展名)"
    except zipfile.BadZipFile:
        print(f"错误: 文件损坏或不是有效的 {ext} (zip) 格式: {path}")
        return 1

    text = (text or "").strip()
    print(f"=== 文件: {path} ===")
    print(f"=== 类型: {ext} | 解析方式: {method} ===")
    print()
    if not text:
        print("(未提取到文本内容)")
        return 0
    if len(text) > args.max_chars:
        print(text[: args.max_chars])
        print(f"[已截断, 原文共 {len(text)} 字符, 仅输出前 {args.max_chars} 字符]")
    else:
        print(text)
    return 0


if __name__ == "__main__":
    sys.exit(main())
